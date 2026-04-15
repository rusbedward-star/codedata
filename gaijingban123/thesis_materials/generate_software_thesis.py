import os
import re
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "基于Python的冰柠牙膏代理销售数据分析.docx"
TARGET = ROOT / "基于Vue与Django的冰柠销量分析与预测系统设计与实现.docx"
MATERIALS_DIR = ROOT / "thesis_materials"
DIAGRAM_SOURCE = MATERIALS_DIR / "系统建模_mermaid.md"
DIAGRAM_DIR = MATERIALS_DIR / "diagrams_generated"
METRICS_FILE = ROOT / "模型评估结果.csv"
FORECAST_FILE = ROOT / "未来12个月预测结果.csv"
MONTHLY_FILE = ROOT / "冰柠销量数据.csv"


TITLE = "基于Vue与Django的冰柠销量分析与预测系统设计与实现"
EN_TITLE = (
    "Design and Implementation of an Agency Sales Analysis and Forecasting "
    "System for a Certain Toothpaste Brand Based on Vue and Django"
)

CHART_IMAGES = {
    "mae": ROOT / "MAE指标对比图.png",
    "rmse": ROOT / "RMSE指标对比图.png",
    "mape": ROOT / "MAPE指标对比图.png",
    "trend": ROOT / "历史与未来销量趋势图.png",
    "multi": ROOT / "未来12个月多模型预测图.png",
    "loss": ROOT / "LSTM训练损失图.png",
}

SYSTEM_SCREENSHOT_DIR = MATERIALS_DIR / "system_screenshots"
SYSTEM_SCREENSHOTS = {
    "dashboard": SYSTEM_SCREENSHOT_DIR / "dashboard.png",
    "data_center": SYSTEM_SCREENSHOT_DIR / "data-center.png",
    "metrics": SYSTEM_SCREENSHOT_DIR / "metrics.png",
    "forecast": SYSTEM_SCREENSHOT_DIR / "forecast.png",
    "charts": SYSTEM_SCREENSHOT_DIR / "charts.png",
    "system": SYSTEM_SCREENSHOT_DIR / "system.png",
}

MERMAID_MARKERS = {
    "business_flow": "## 1. 业务流程图",
    "use_case": "## 2. 系统用例图",
    "class_diagram": "## 3. 系统类图",
    "sequence_diagram": "## 4. 系统顺序图",
    "activity_diagram": "## 5. 系统活动图",
    "er_diagram": "## 6. ER 图",
}

REFERENCES = [
    "[1] 王家聚. 基于B/S的房地产销售系统的设计与实现[J]. 海南大学学报(自然科学版), 2009, 27(1): 60-64.",
    "[2] 徐燕萍. 基于B/S的企业销售管理系统设计[J]. 软件, 2016, 37(5): 84-88.",
    "[3] 王思辰, 李林. 基于Vue.js的电商管理平台的设计与实现[J]. 现代信息科技, 2021, 5(14): 13-16.",
    "[4] 刘正. 基于Vue的地下综合管廊管理平台的前端设计与实现[J]. 现代信息科技, 2021, 5(16): 13-18.",
    "[5] 龙涛, 谌爱文. 基于SaaS模式的云农平台的设计与搭建[J]. 现代信息科技, 2021, 5(14): 1-4.",
    "[6] 徐秀芳, 夏旻, 徐森, 等. 基于Django的校园疫情防控系统设计与实现[J]. 软件导刊, 2021, 20(2): 24-30.",
    "[7] 赵宸, 刘建华. 基于Django的分布式爬虫框架设计与实现[J]. 计算机与数字工程, 2020, 48(10): 2495-2498.",
    "[8] 林晨, 邓录, 董璐. 基于SVG和Vue的数据过程可视化设计[J]. 计算机系统应用, 2022, 31(4): 130-136.",
    "[9] 方生. 基于“Vue.js”前端框架技术的研究[J]. 电脑知识与技术, 2021, 17(19): 59-60,64.",
    "[10] 李嘉, 赵凯强, 李长云. Web前端开发技术的演化与MVVM设计模式研究[J]. 电脑知识与技术, 2018, 14(2): 221-222,251.",
    "[11] 骆海东, 马卫清, 梁丹. 面向零售电商的仓库管理系统设计[J]. 现代信息科技, 2021, 5(2): 27-31.",
    "[12] 陈宇. 基于Spring Boot的电商管理系统的设计[J]. 现代信息科技, 2020, 4(1): 25-26.",
    "[13] 杨章伟, 肖异骐. 基于SSM+Vue的赣西傩资源数字化系统设计与实现[J]. 网络安全技术与应用, 2022(4): 44-45.",
    "[14] 张倩, 李旭英, 林华焜, 等. 基于Vue.js+Koa框架的APP平台设计与实现:以酒类文化交流与电子商务为例[J]. 现代商业, 2021(8): 76-78.",
    "[15] 闫锦彪, 杨冬梅, 张进. 校园二手教材网络交易平台的构建方法研究[J]. 现代信息科技, 2021, 5(2): 107-110.",
    "[16] Fielding R T, Taylor R N. Principled Design of the Modern Web Architecture[J]. ACM Transactions on Internet Technology, 2002, 2(2): 115-150.",
    "[17] Ceri S, Fraternali P, Bongio A. Web Modeling Language (WebML): A Modeling Language for Designing Web Sites[J]. Computer Networks, 2000, 33(1-6): 137-157.",
    "[18] Deshpande Y, Murugesan S, Ginige A, et al. Web Engineering[J]. Journal of Web Engineering, 2002, 1(1): 3-17.",
    "[19] Hochreiter S, Schmidhuber J. Long Short-Term Memory[J]. Neural Computation, 1997, 9(8): 1735-1780.",
    "[20] Paszke A, Gross S, Massa F, et al. PyTorch: An Imperative Style, High-Performance Deep Learning Library[C]//Advances in Neural Information Processing Systems 32. Vancouver: Curran Associates, 2019: 8024-8035.",
]


def set_run_font(run, size=12, bold=False):
    font = run.font
    font.name = "宋体"
    font.size = Pt(size)
    font.bold = bold
    rpr = run._element.get_or_add_rPr()
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rpr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")


def set_paragraph_text(paragraph, text, size=12, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def format_paragraph(paragraph, role="body"):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.5
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if role == "heading1":
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(6)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif role == "heading2":
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(4)
        fmt.space_after = Pt(4)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif role == "heading3":
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(2)
        fmt.space_after = Pt(2)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif role == "caption":
        fmt.first_line_indent = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif role == "body":
        fmt.first_line_indent = Cm(0.74)


def add_paragraph(doc, text, style=None, role="body", size=12, bold=False):
    try:
        paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    except KeyError:
        paragraph = doc.add_paragraph()
    set_paragraph_text(paragraph, text, size=size, bold=bold)
    format_paragraph(paragraph, role=role)
    return paragraph


def add_heading(doc, text, level):
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
    role_map = {1: "heading1", 2: "heading2", 3: "heading3"}
    size_map = {1: 16, 2: 14, 3: 12}
    return add_paragraph(
        doc,
        text,
        style=style_map[level],
        role=role_map[level],
        size=size_map[level],
        bold=True,
    )


def add_body(doc, *paragraphs):
    for index, paragraph in enumerate(paragraphs):
        role = "body"
        style = "First Paragraph" if index == 0 else "Body Text"
        add_paragraph(doc, paragraph, style=style, role=role)


def add_table_caption(doc, text):
    add_paragraph(doc, text, style="图表名", role="caption")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        if edge_data:
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
        else:
            element.set(qn("w:val"), "nil")


def make_three_line_table(table):
    rows = table.rows
    cols = len(table.columns)
    for row in rows:
        for cell in row.cells:
            set_cell_border(cell, top=None, bottom=None, left=None, right=None)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if paragraph.runs:
                    for run in paragraph.runs:
                        set_run_font(run, size=10.5)

    for col in range(cols):
        header_cell = rows[0].cells[col]
        set_cell_border(
            header_cell,
            top={"val": "single", "sz": 8, "space": 0, "color": "000000"},
            bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
            left=None,
            right=None,
        )
        for paragraph in header_cell.paragraphs:
            if paragraph.runs:
                for run in paragraph.runs:
                    set_run_font(run, size=10.5, bold=True)

        footer_cell = rows[-1].cells[col]
        set_cell_border(
            footer_cell,
            top=None,
            bottom={"val": "single", "sz": 8, "space": 0, "color": "000000"},
            left=None,
            right=None,
        )


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(value)
    make_three_line_table(table)
    return table


def make_borderless_table(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=None, bottom=None, left=None, right=None)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if paragraph.runs:
                    for run in paragraph.runs:
                        set_run_font(run, size=12)


def insert_table_after(anchor_table, table):
    anchor_table._element.addnext(table._element)


def insert_cover_table(doc):
    cover_rows = [
        ("学    院", "计算机学院"),
        ("专    业", "计算机科学与技术"),
        ("班    级", ""),
        ("学    号", ""),
        ("学生姓名", ""),
        ("指导教师", ""),
    ]

    if len(doc.tables) > 1 and doc.tables[1].cell(0, 0).text.strip() == "学    院":
        table = doc.tables[1]
    else:
        table = doc.add_table(rows=len(cover_rows), cols=2)
        insert_table_after(doc.tables[0], table)

    for row_index, (left, right) in enumerate(cover_rows):
        table.cell(row_index, 0).text = left
        table.cell(row_index, 1).text = right
    make_borderless_table(table)


def normalize_caption_numbers(doc):
    pattern = re.compile(r"([图表])(\d+)\.(\d+)")

    def replace_text(text):
        return pattern.sub(r"\1\2-\3", text)

    for paragraph in doc.paragraphs:
        if paragraph.text:
            new_text = replace_text(paragraph.text)
            if new_text != paragraph.text:
                set_paragraph_text(paragraph, new_text, size=12, bold=False)
                format_paragraph(
                    paragraph,
                    role="caption" if new_text.startswith(("图", "表")) else "body",
                )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        new_text = replace_text(paragraph.text)
                        if new_text != paragraph.text:
                            set_paragraph_text(paragraph, new_text, size=10.5, bold=False)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_figure(doc, image_path, caption, width_cm=15.5):
    if not image_path.exists():
        add_paragraph(doc, f"{caption}图片资源缺失，生成文档时未插入。", style="Body Text")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    add_table_caption(doc, caption)


def remove_body_from_paragraph(doc, paragraph_index):
    start_element = doc.paragraphs[paragraph_index]._element
    body = doc.element.body
    children = list(body)
    start_idx = children.index(start_element)
    for element in children[start_idx:]:
        body.remove(element)


def extract_mermaid_block(text, marker):
    section = text.split(marker, 1)[1]
    match = re.search(r"```mermaid\n(.*?)```", section, flags=re.S)
    return match.group(1).strip()


def render_mermaid_diagrams():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    text = DIAGRAM_SOURCE.read_text(encoding="utf-8")
    env = os.environ.copy()
    env["npm_config_cache"] = str(ROOT / ".npm-cache")
    outputs = {}

    for key, marker in MERMAID_MARKERS.items():
        code = extract_mermaid_block(text, marker)
        mmd_path = DIAGRAM_DIR / f"{key}.mmd"
        png_path = DIAGRAM_DIR / f"{key}.png"
        mmd_path.write_text(code, encoding="utf-8")
        subprocess.run(
            [
                "npx",
                "-y",
                "@mermaid-js/mermaid-cli",
                "-i",
                str(mmd_path),
                "-o",
                str(png_path),
                "-b",
                "white",
            ],
            cwd=ROOT,
            env=env,
            check=True,
        )
        outputs[key] = png_path
    return outputs


def build_stats():
    monthly_df = pd.read_csv(MONTHLY_FILE)
    metrics_df = pd.read_csv(METRICS_FILE)
    forecast_df = pd.read_csv(FORECAST_FILE)

    best_rmse = metrics_df.sort_values("RMSE").iloc[0]
    best_mape = metrics_df.sort_values("MAPE(%)").iloc[0]
    peak_value = None
    peak_month = None
    peak_model = None
    for _, row in forecast_df.iterrows():
        for column in forecast_df.columns:
            if column == "月份":
                continue
            value = float(row[column])
            if peak_value is None or value > peak_value:
                peak_value = value
                peak_month = row["月份"]
                peak_model = column

    return {
        "monthly_df": monthly_df,
        "metrics_df": metrics_df,
        "forecast_df": forecast_df,
        "sample_count": len(monthly_df),
        "date_start": monthly_df["date"].min(),
        "date_end": monthly_df["date"].max(),
        "avg_sales": round(float(monthly_df["sales"].mean()), 2),
        "max_sales": round(float(monthly_df["sales"].max()), 2),
        "min_sales": round(float(monthly_df["sales"].min()), 2),
        "best_rmse_name": best_rmse["模型"],
        "best_rmse": round(float(best_rmse["RMSE"]), 4),
        "best_mape_name": best_mape["模型"],
        "best_mape": round(float(best_mape["MAPE(%)"]), 4),
        "peak_month": peak_month,
        "peak_model": peak_model,
        "peak_value": round(float(peak_value), 4),
    }


def update_front_matter(doc):
    doc.tables[0].cell(0, 0).text = TITLE
    insert_cover_table(doc)

    if len(doc.paragraphs) > 5:
        set_paragraph_text(doc.paragraphs[5], "二〇二六 年 四 月", size=12)
        doc.paragraphs[5].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_paragraph_text(doc.paragraphs[27], "摘    要", size=16, bold=True)
    set_paragraph_text(
        doc.paragraphs[28],
        (
            "围绕代理销售研究中偏重数据分析、软件成果支撑不足这一问题，本文完成了一套基于Vue、"
            "Django、PyTorch与Pyecharts的冰柠销量分析与预测系统。系统采用前后端分离"
            "架构，包含数据展示、模型评估、结果查询和动态图表等功能，同时完成了需求分析、系统设"
            "计、功能实现和单元测试。结果表明，该系统能够较稳定地完成销售数据展示、模型切换、预"
            "测查询与图表联动，整体上具有较好的可用性和扩展性。"
        ),
        size=12,
    )
    set_paragraph_text(
        doc.paragraphs[30],
        "关键词：冰柠；代理销售系统；前后端分离；销量预测；软件设计",
        size=12,
    )
    set_paragraph_text(doc.paragraphs[32], EN_TITLE, size=14, bold=True)
    set_paragraph_text(
        doc.paragraphs[33],
        (
            "This paper develops an agency sales analysis and forecasting system for a certain "
            "toothpaste brand based on Vue, Django, PyTorch and Pyecharts. The system uses a "
            "front-end and back-end separation architecture and integrates data display, model "
            "evaluation, forecast querying and dynamic visualization. The thesis mainly covers "
            "requirement analysis, system design, implementation and unit testing. The results "
            "show that the system can support sales data display, model switching, forecast "
            "querying and chart interaction in a stable manner."
        ),
        size=11,
    )
    set_paragraph_text(
        doc.paragraphs[35],
        "Keywords: toothpaste brand; sales system; Vue; Django; forecasting; software design",
        size=11,
    )

    toc_lines = [
        "第1章 绪论\t1",
        "1.1 课题背景与研究意义\t1",
        "1.1.1 研究背景\t1",
        "1.1.2 研究意义\t2",
        "1.2 国内外研究现状\t3",
        "1.3 研究内容与技术路线\t4",
        "1.3.1 研究内容\t4",
        "1.3.2 技术路线\t5",
        "1.4 论文结构安排\t6",
        "第2章 系统分析\t7",
        "2.1 需求分析\t7",
        "2.1.1 功能需求\t7",
        "2.1.2 数据需求\t8",
        "2.1.3 非功能需求\t9",
        "2.2 业务流程分析\t10",
        "2.2.1 销量数据管理流程\t10",
        "2.2.2 模型评估与预测流程\t11",
        "2.2.3 图表可视化流程\t12",
        "2.3 用例分析与规约\t13",
        "2.4 本章小结\t15",
        "第3章 系统设计\t16",
        "3.1 总体架构设计\t16",
        "3.2 功能模块设计\t18",
        "3.3 系统建模设计\t20",
        "3.4 数据库设计\t23",
        "3.5 安全与性能设计\t25",
        "3.6 本章小结\t26",
        "第4章 系统实现\t27",
        "4.1 开发环境与工程结构\t27",
        "4.2 前端功能实现\t28",
        "4.3 后端接口实现\t31",
        "4.4 预测与可视化实现\t34",
        "4.5 本章小结\t38",
        "第5章 系统测试\t39",
        "5.1 测试环境与策略\t39",
        "5.2 模块单元测试\t40",
        "5.3 接口联调与结果分析\t44",
        "5.4 本章小结\t46",
        "第6章 总结与展望\t47",
        "参考文献\t49",
        "致谢\t51",
    ]
    start = 38
    for idx, line in enumerate(toc_lines):
        set_paragraph_text(doc.paragraphs[start + idx], line, size=12)
    for idx in range(start + len(toc_lines), 82):
        set_paragraph_text(doc.paragraphs[idx], "", size=12)


def chapter_1(doc, stats):
    add_heading(doc, "第1章 绪论", 1)
    add_heading(doc, "1.1 课题背景与研究意义", 2)
    add_heading(doc, "1.1.1 研究背景", 3)
    add_body(
        doc,
        (
            "在快速消费品代理销售场景中，业务数据往往同时承担经营复盘、库存调度、区域投放和"
            "营销评估等多种作用。传统做法通常依赖Excel表格、静态统计图和人工经验完成分析，虽"
            "然能够支持基础汇总，但在模型切换、结果联动和图形化展示方面存在明显不足。对于毕业设"
            "计而言，如果论文仅停留在算法对比与数据分析层面，难以完整体现需求分析、系统设计、功"
            "能实现和测试验证等软件工程能力，因此需要将原有研究进一步扩展为一个可运行的软件系统。"
        ),
        (
            f"本文所研究的业务对象为冰柠代理销售场景，现有样本覆盖{stats['date_start']}至"
            f"{stats['date_end']}共{stats['sample_count']}条月度记录，平均销量为{stats['avg_sales']}，"
            f"峰值销量达到{stats['max_sales']}。这些真实业务数据为系统中的数据中心、模型评估、未来"
            "预测和图表展示提供了基础支撑。基于这些数据，论文不再以“分析某品牌销量规律”为唯一"
            "目标，而是转向“设计并实现一套销售分析与预测系统”，从题目、章节组织到成果形式全部与"
            "软件设计类毕业论文要求保持一致。"
        ),
        (
            "从技术实现角度看，Vue在前端页面组织、响应式渲染和路由管理方面具有清晰的组件化优"
            "势，Django在接口组织、数据封装和服务端渲染图表方面具备较成熟的工程能力，PyTorch适"
            "合封装深度学习预测模型，Pyecharts则便于快速构建适合论文展示的动态图表。将这些技术"
            "组合到同一系统中，既能保留原始销售预测研究中的算法价值，又能补足软件设计与实现层面"
            "的毕业论文要求[1-3]。"
        ),
    )
    add_heading(doc, "1.1.2 研究意义", 3)
    add_body(
        doc,
        (
            "本课题的意义可以从几个层面来理解。其一，本文以代理销售业务为背景，把需求分析、系统"
            "建模、前后端分离架构、数据库设计和预测模型集成放在同一研究框架下展开，这样处理之后，"
            "从数据处理到软件交付之间的实现路径会更完整一些。其二，从工程实现看，系统借助路由页"
            "面、REST接口、结果聚合服务和动态图表服务，把页面层和模型层拆开了，后面如果要继续补"
            "新模型、新图表或新模块，修改起来也会相对方便。其三，从实际应用看，系统可以为代理销"
            "售场景中的销量预测、库存预判、促销复盘和答辩展示提供更直观的支持。"
        ),
        (
            "与原有题目相比，本文的重点不再是“用什么模型预测更好”，而是“如何将数据分析结果转"
            "化为可使用、可测试、可演示的系统功能”。在毕业论文评价标准中，这种转向能够更好地体"
            "现软件设计类选题所要求的体系结构设计、模块划分、顺序分析、数据库设计和单元测试能力。"
            "同时，系统中所有涉及品牌名称的展示均已统一替换为“冰柠”，满足论文规范化写作"
            "和敏感信息处理要求。"
        ),
    )

    add_heading(doc, "1.2 国内外研究现状", 2)
    add_body(
        doc,
        (
            "国内相关研究主要集中在两个方向。其一是销售管理系统、B/S业务系统及其数据库设计，"
            "如基于B/S模式的销售系统、企业管理系统和电商平台管理系统研究，为本课题提供了系统模"
            "块划分、业务流程抽象和测试方法方面的参考[4-5]。其二是基于Vue、Django等现代Web框"
            "架的信息平台设计与实现研究，相关成果表明，组件化前端与轻量级后端结合能够较好地支持"
            "数据展示、查询交互与可视化服务[6-13]。此外，关于SaaS平台、管理信息系统和多租户数据"
            "结构的研究，也为本文在扩展性和接口解耦方面提供了设计依据[3][14-15]。"
        ),
        (
            "国外研究则更多从Web体系结构、Web建模方法和智能预测模型应用角度展开。Fielding与"
            "Taylor从现代Web架构原则出发，对资源、统一接口和系统可扩展性进行了系统讨论[16]；"
            "Ceri等提出的WebML方法则为Web应用的页面组织、数据视图和导航关系建模提供了清晰路径"
            "[17]；Deshpande等对Web Engineering的研究进一步强调了Web系统在需求、设计、实现和测"
            "试上的工程化方法[18]。在预测模型方面，Hochreiter和Schmidhuber提出的LSTM为时间序列"
            "预测提供了深度学习基础，Paszke等则通过PyTorch工作证明了动态图框架在模型构建与部署"
            "中的工程优势[19-20]。"
        ),
        (
            "可以看出，现有研究在“系统实现”和“销售预测”之间还是有些分开的。系统类论文往往更关"
            "注业务流程、页面组织和数据库设计，对预测模型切换和动态图表联动涉及得不算多；算法类"
            "论文则更侧重误差指标和模型训练过程，但软件交付部分通常比较弱。本文并没有把重点放在"
            "提出新的预测算法上，而是希望把已有分析结果转成可运行、可展示、也便于后续维护的软件"
            "系统。"
        ),
    )

    add_heading(doc, "1.3 研究内容与技术路线", 2)
    add_heading(doc, "1.3.1 研究内容", 3)
    add_body(
        doc,
        (
            "本文的研究内容基本围绕“系统分析、系统设计、系统实现、系统测试”这几部分展开。先对"
            "业务需求做梳理，把数据查看、模型比较、未来预测查询、图表可视化和系统设计展示这些核"
            "心功能明确下来；再完成系统总体架构、前后端模块、接口结构、类关系和数据库逻辑模型设"
            "计；之后基于Vue与Django实现多页面前端和统一接口后端，并把随机森林、梯度提升回归、"
            "极端随机森林、线性回归和LSTM的结果统一封装起来；最后再对页面模块、接口模块和图表模"
            "块做单元测试与联调测试，验证系统能否正常运行。"
        ),
        (
            "为了满足软件设计类毕业论文模板，本文除了实现系统本身，还补充了业务流程图、用例图、"
            "类图、顺序图、活动图、ER图和用例规约表等建模成果，并在第五章中对首页模块、数据中"
            "心模块、模型评估模块、预测结果模块、图表中心模块以及后端接口模块分别进行测试说明。"
            "这样既保证论文正文充实，也使系统成果与论文章节紧密对应。"
        ),
    )
    add_heading(doc, "1.3.2 技术路线", 3)
    add_body(
        doc,
        (
            "系统技术路线可概括为“数据准备、结果生成、接口封装、页面展示、测试验证”五个阶段。"
            "首先对真实销售数据进行清洗、字段整理和时间序列组织；其次离线生成模型评估结果与未来"
            "12个月预测结果；然后使用Django封装概览、样本、指标、预测、详情和图表接口；在前端"
            "通过Vue Router组织页面跳转，并使用统一数据加载组合式函数完成多页面共享；最后通过"
            "Django TestCase和模块测试表验证系统功能。"
        ),
        (
            "从工程实现角度看，系统采用前后端分离模式。前端负责导航、表格渲染、模型选择和图表"
            "嵌入，后端负责结果读取、业务整合和Pyecharts图表渲染。为了降低数据依赖复杂度，当前"
            "系统以CSV结果文件作为核心数据源，同时在论文中给出面向数据库落地的逻辑设计，为后续"
            "接入SQLite或MySQL提供扩展空间。该路线既能够快速形成可运行成果，又保留了向更完整管"
            "理系统演进的工程基础。"
        ),
    )

    add_heading(doc, "1.4 论文结构安排", 2)
    add_body(
        doc,
        (
            "全文共分为六章。第一章为绪论，说明课题背景、研究意义、国内外研究现状及技术路线。"
            "第二章为系统分析，重点讨论需求、业务流程和用例规约。第三章为系统设计，给出总体架构、"
            "模块设计、系统建模、数据库设计和性能安全设计。第四章为系统实现，介绍前端页面、后端"
            "接口、预测模型封装和可视化实现。第五章为系统测试，对各功能模块的单元测试和联调测试"
            "进行分析。第六章为总结与展望，对本文工作进行归纳并提出后续改进方向。"
        ),
    )


def chapter_2(doc, stats, diagrams):
    add_heading(doc, "第2章 系统分析", 1)
    add_heading(doc, "2.1 需求分析", 2)
    add_heading(doc, "2.1.1 功能需求", 3)
    add_body(
        doc,
        (
            "系统的主要使用者包括论文答辩演示者、课程指导教师和业务分析人员。对这几类用户来说，"
            "比较直接的需求还是尽快看到真实销售样本、模型评估指标、不同模型切换后的结果以及图形"
            "化展示内容。也就是说，系统至少需要具备首页概览、数据中心、模型评估、预测结果、图表"
            "中心和系统设计这六个页面，而且页面之间要能顺畅切换。首页用于展示系统规模、模型数量"
            "和推荐模型，数据中心展示样本时间范围、技术栈和原始数据预览，模型评估页展示MAE、"
            "RMSE、MAPE指标，预测结果页按模型和月份查看未来值，图表中心集中展示动态图表，系统设"
            "计页则用于呈现模块设计与优化信息。"
        ),
        (
            "从功能边界来看，当前系统虽然还没有加入用户登录、权限细分和在线训练这些内容，但基本"
            "的软件闭环已经形成了。系统能够围绕同一批业务数据给出多模型结果对比，也能通过统一接"
            "口向前端输出结构化数据，再配合动态图表完成展示。换言之，它已经能够满足毕业设计演示"
            "场景的主要需求，同时也给后续补充角色权限、模型重训和日志记录等功能留出了空间。"
        ),
    )
    add_table_caption(doc, "表2.1 系统功能需求分析表")
    add_table(
        doc,
        [
            ["功能模块", "输入", "处理逻辑", "输出结果"],
            ["系统首页", "概览接口请求", "汇总样本量、模型数与推荐模型", "统计卡片、快捷入口"],
            ["数据中心", "样本数据接口请求", "读取真实销售数据并格式化", "月度样本表、时间范围"],
            ["模型评估", "评估接口请求", "读取指标结果并按误差排序", "模型评估表、推荐模型"],
            ["预测结果", "预测与详情请求", "按模型加载未来12个月预测值", "预测表、重点模型预览"],
            ["图表中心", "图表接口请求", "生成指标图、趋势图与模型预测图", "Pyecharts动态图"],
            ["系统设计", "模块接口请求", "返回前后端模块与优化项", "模块清单、优化过程"],
        ],
    )

    add_heading(doc, "2.1.2 数据需求", 3)
    add_body(
        doc,
        (
            "系统的数据需求主要包括三类。第一类是业务样本数据，即销售月份、销量值、上月销量、"
            "去年同月销量、月份编号、是否节假日、是否促销、区域和产品系列等字段。这些数据用于支"
            "撑数据中心展示，同时也是原始模型训练与特征构造的重要来源。第二类是模型评估数据，即"
            "各模型的MAE、RMSE和MAPE指标结果，用于支持模型评估页面的横向比较与推荐模型判定。"
            "第三类是未来预测数据，即各模型对未来12个月销量的预测结果，用于支持预测结果页面和图"
            "表中心中的多模型对比。"
        ),
        (
            "数据需求不仅体现在字段完备性上，还体现在数据一致性上。系统要求模型名称在评估文件、"
            "预测文件和图表接口之间保持一致，否则会导致模型切换时无法完成联动。系统还要求月份字"
            "段在前端能够直接用于模糊搜索和折线图横轴映射，因此采用统一的“YYYY-MM”格式。对于"
            "论文写作而言，这种数据需求分析能够直接对应数据库设计章节中的实体字段说明。"
        ),
    )

    add_heading(doc, "2.1.3 非功能需求", 3)
    add_body(
        doc,
        (
            "在非功能需求方面，首先要求系统具备良好的可用性。由于系统主要用于答辩与演示，页面风"
            "格不需要过度装饰，但必须保持结构清晰、响应及时、跳转明确。其次要求系统具备稳定性，"
            "即在结果文件存在的情况下，所有核心接口均能够返回合法JSON或HTML图表结果。再次要求"
            "系统具备可扩展性，后续能够平滑接入数据库持久化、用户认证、模型重训、上传导入和日志"
            "监控等扩展能力。最后要求系统具备一定可维护性，通过前后端分离架构降低模块间耦合度。"
        ),
        (
            "此外，论文场景下还存在展示性要求。系统界面需适合投影展示，表格和图表应具有较高可读"
            "性，图像与模块说明应能够被直接截图用于论文插图。这意味着系统除满足业务逻辑外，还需"
            "在布局、字体、配色和信息密度上兼顾学术展示需求，这也是本课题在需求分析阶段就必须考"
            "虑的重要约束。"
        ),
    )

    add_heading(doc, "2.2 业务流程分析", 2)
    add_heading(doc, "2.2.1 销量数据管理流程", 3)
    add_body(
        doc,
        (
            "销量数据管理流程主要发生在数据中心页面。用户进入页面后，前端通过统一数据加载逻辑请"
            "求样本数据接口；后端读取真实销售CSV文件，对字段名和空值进行标准化处理后返回结构化"
            "结果；前端再将结果渲染为表格，供用户查看月份、销量、区域和产品系列等信息。该流程虽"
            "然处理逻辑相对简单，但在系统中承担着“数据来源可信展示”的作用，是整套论文能够说明使"
            "用真实业务数据的重要基础。"
        ),
    )
    add_heading(doc, "2.2.2 模型评估与预测流程", 3)
    add_body(
        doc,
        (
            "模型评估与预测流程是系统的核心业务流程之一。用户在模型评估页面查看整体指标时，系统"
            "首先加载评估结果并根据RMSE值确定默认推荐模型；当用户进入预测结果页面并选择某一模型"
            "后，页面会继续请求模型详情接口，后端据此从评估文件和预测文件中提取对应模型的指标与"
            "未来预测序列。通过这种“总览先行、详情按需加载”的方式，系统既保证了页面首次加载效率，"
            "又实现了多模型之间的灵活切换。"
        ),
        (
            f"以当前结果为例，系统默认推荐{stats['best_rmse_name']}模型，因为其RMSE为"
            f"{stats['best_rmse']}，在现有模型中表现最好；而在MAPE指标上，{stats['best_mape_name']}"
            f"模型更优，说明不同指标对应的模型选择结论可能并不完全一致。系统通过展示这种差异，帮"
            "助用户理解“评价指标选择会影响推荐模型”的业务含义，从而提高系统的解释性。"
        ),
    )
    add_heading(doc, "2.2.3 图表可视化流程", 3)
    add_body(
        doc,
        (
            "图表可视化流程发生在图表中心页面。页面加载时首先读取预测数据的列名，生成模型下拉列"
            "表；之后根据页面参数判断应展示指标类图表、趋势类图表还是全部图表；如果用户切换模型，"
            "则页面会重新构造模型预测图的请求地址，从而使iframe中嵌入的Pyecharts图实时刷新。后"
            "端图表接口则根据图表类型调用不同的构造函数，动态返回柱状图或折线图HTML。"
        ),
        (
            "这一流程的价值在于打通了页面交互与服务端图表渲染。与直接在前端使用静态图片不同，本"
            "系统将图表构建逻辑保留在后端，使图表数据与模型结果文件保持同步，也方便后续接入新的"
            "图表类型和查询参数。"
        ),
    )
    add_figure(doc, diagrams["business_flow"], "图2.1 系统业务流程图", width_cm=16)

    add_heading(doc, "2.3 用例分析与规约", 2)
    add_body(
        doc,
        (
            "根据需求分析，系统参与者可以抽象为“系统用户”一个角色。该角色可以完成查看系统首页、"
            "查看样本数据、查看模型评估结果、切换预测模型、查看未来预测结果、查看动态图表和查看"
            "系统设计信息等操作。系统用例图从宏观上反映了用户与系统功能之间的交互边界，为后续界"
            "面设计和测试用例设计提供了结构化依据。"
        ),
    )
    add_figure(doc, diagrams["use_case"], "图2.2 系统用例图", width_cm=15.5)

    add_table_caption(doc, "表2.2 查看模型评估结果用例规约表")
    add_table(
        doc,
        [
            ["项目", "内容"],
            ["用例编号", "UC-01"],
            ["用例名称", "查看模型评估结果"],
            ["执行者", "系统用户"],
            ["前置条件", "系统正常启动，评估结果文件存在"],
            ["后置条件", "页面成功显示模型指标表并标识推荐模型"],
            ["基本事件流", "进入评估页面→调用/api/metrics/→读取结果→返回表格→高亮推荐模型"],
            ["备选事件流", "结果文件缺失或接口异常时，页面显示错误提示"],
            ["业务规则", "默认依据RMSE最小值确定推荐模型"],
        ],
    )

    add_table_caption(doc, "表2.3 查看未来预测结果用例规约表")
    add_table(
        doc,
        [
            ["项目", "内容"],
            ["用例编号", "UC-02"],
            ["用例名称", "查看未来预测结果"],
            ["执行者", "系统用户"],
            ["前置条件", "预测结果文件存在，前端页面正常访问"],
            ["后置条件", "页面显示未来12个月多模型预测表"],
            ["基本事件流", "进入预测页→请求/api/forecast/→按月份筛选→展示预测表"],
            ["备选事件流", "预测文件不存在时返回空表或错误信息"],
            ["业务规则", "月份筛选只影响表格展示，不改变原始数据源"],
        ],
    )

    add_table_caption(doc, "表2.4 查看指定模型预测图用例规约表")
    add_table(
        doc,
        [
            ["项目", "内容"],
            ["用例编号", "UC-03"],
            ["用例名称", "查看指定模型预测图"],
            ["执行者", "系统用户"],
            ["前置条件", "图表接口可访问，所选模型存在"],
            ["后置条件", "页面显示所选模型的预测折线图"],
            ["基本事件流", "选择模型→请求/api/model-detail/→请求/api/charts/model-forecast/→渲染图表"],
            ["备选事件流", "模型名称不存在时返回404并提示加载失败"],
            ["业务规则", "切换模型后图表需立即刷新，且指标信息同步更新"],
        ],
    )

    add_heading(doc, "2.4 本章小结", 2)
    add_body(
        doc,
        (
            "本章从功能需求、数据需求、非功能需求和业务流程等角度对系统进行了系统分析，并给出"
            "了用例图与用例规约表。分析结果表明，系统的核心价值在于把真实销售数据、多模型评估结"
            "果与动态图表服务统一到同一业务平台中，为后续架构设计和模块实现提供了明确边界。"
        ),
    )


def chapter_3(doc, stats, diagrams):
    add_heading(doc, "第3章 系统设计", 1)
    add_heading(doc, "3.1 总体架构设计", 2)
    add_heading(doc, "3.1.1 设计原则", 3)
    add_body(
        doc,
        (
            "系统设计时主要考虑了职责分离、模块内聚、接口统一和后续扩展这几个方面。前端更多负责"
            "导航、交互和视图组织，后端则负责数据读取、结果封装和图表渲染；同一模块内部尽量只处"
            "理同类业务；数据接口保持相近的JSON输出格式；如果后面还要补页面、图表或预测模型，也"
            "尽量不要大改现有结构。这样处理的好处比较直接，一方面能保证答辩演示时系统运行稳定，"
            "另一方面也给后续继续完善业务平台留下了余地。"
        ),
    )
    add_heading(doc, "3.1.2 技术架构", 3)
    add_body(
        doc,
        (
            "系统总体采用前后端分离架构。前端基于Vue 3、Vite和Vue Router实现单页应用式的多视"
            "图导航，使用组合式函数统一加载系统概览、样本数据、评估结果、预测结果和系统模块说明；"
            "后端基于Django实现URL分发、控制器逻辑和图表服务，将真实销售数据、评估结果文件、预测"
            "结果文件以及图像资源统一封装为标准接口；算法结果层通过Python脚本离线生成CSV文件；"
            "可视化层通过Pyecharts构建服务端图表。该架构符合现代Web系统设计的主流模式[6-11][16-18]。"
        ),
    )
    add_heading(doc, "3.1.3 部署结构", 3)
    add_body(
        doc,
        (
            "在开发阶段，系统采用本地分离部署方式。前端通过Vite开发服务器运行，并将/api请求代理"
            "至Django服务；后端通过Django开发服务器提供JSON接口与图表HTML页面。部署结构虽然简化，"
            "但已完整覆盖浏览器端、服务端和数据文件层三个层次。若后续需要上线部署，可进一步采用"
            "Nginx反向代理、Gunicorn或uWSGI应用容器以及MySQL数据库进行生产化扩展。"
        ),
    )

    add_heading(doc, "3.2 功能模块设计", 2)
    add_heading(doc, "3.2.1 前端模块设计", 3)
    add_body(
        doc,
        (
            "前端模块按照页面职责划分为六个视图模块和一个公共组件模块。DashboardView负责展示系"
            "统总体概览和快捷入口；DataCenterView负责展示样本数据和时间范围；MetricsView负责展"
            "示评估结果及推荐模型；ForecastView负责月份筛选、重点模型选择和局部预览；ChartsView"
            "负责动态图表嵌入与模型切换；SystemView负责展示前后端模块和系统优化过程。公共的"
            "DataTable组件用于统一渲染表格，提高页面之间的一致性。"
        ),
        (
            "此外，useSystemData组合式函数承担前端数据聚合职责。该函数在内部并发请求概览、评估、"
            "预测、图片、模块说明、样本和系统建议接口，将结果缓存到响应式状态中，从而避免各页面"
            "重复编写相同的数据加载逻辑。这种做法既减少了代码冗余，也让页面层代码保持简洁，更符"
            "合毕业设计中模块复用和解耦的设计要求。"
        ),
    )

    add_heading(doc, "3.2.2 后端接口设计", 3)
    add_body(
        doc,
        (
            "后端接口设计采用面向资源的组织方式。概览接口负责返回项目名称、技术栈、样本规模、图"
            "表数量和推荐模型；评估接口返回模型指标列表；预测接口返回未来预测表；样本数据接口返"
            "回真实业务数据预览；模型详情接口返回指定模型的指标与时间序列；图表接口返回动态图表"
            "HTML；图片接口返回静态结果图；模块接口返回前后端模块说明和优化项。接口之间虽然功能"
            "不同，但返回结构整体统一，便于前端按模块调用。"
        ),
    )
    add_table_caption(doc, "表3.1 核心接口设计表")
    add_table(
        doc,
        [
            ["接口路径", "请求方式", "功能说明", "返回类型"],
            ["/api/overview/", "GET", "获取系统概览和推荐模型", "JSON"],
            ["/api/sample-data/", "GET", "获取真实销售样本预览", "JSON"],
            ["/api/metrics/", "GET", "获取模型评估结果", "JSON"],
            ["/api/forecast/", "GET", "获取未来12个月预测结果", "JSON"],
            ["/api/model-detail/?model=*", "GET", "获取指定模型详情与序列", "JSON"],
            ["/api/charts/<chart_key>/", "GET", "获取指标图或趋势图", "HTML"],
            ["/api/charts/model-forecast/?model=*", "GET", "获取指定模型预测图", "HTML"],
            ["/api/modules/", "GET", "获取系统模块设计说明", "JSON"],
        ],
    )

    add_heading(doc, "3.2.3 模型服务设计", 3)
    add_body(
        doc,
        (
            "模型服务设计强调结果统一封装而非在线训练。随机森林、梯度提升回归、极端随机森林、"
            "线性回归和LSTM等模型的评估结果均被整理为统一CSV格式，模型详情接口能够根据模型名从"
            "指标文件和预测文件中提取同一模型的结果，前端无需关心模型训练细节即可完成切换展示。"
            "这种设计特别适合论文答辩场景，因为它既保留了模型多样性，又避免在线训练带来的时延和"
            "运行不确定性。"
        ),
    )
    add_heading(doc, "3.2.4 可视化服务设计", 3)
    add_body(
        doc,
        (
            "可视化服务由Pyecharts图表构建函数组成，根据请求路径区分为指标柱状图、历史趋势图、多"
            "模型预测折线图和指定模型预测折线图四类。图表服务采用服务端生成HTML的方式返回，前端"
            "只需要通过iframe进行嵌入。这样做的优点在于图表生成逻辑统一集中于后端，且与真实数据"
            "源保持一致，减少了前端重复拼装ECharts配置的复杂度。"
        ),
    )

    add_heading(doc, "3.3 系统建模设计", 2)
    add_heading(doc, "3.3.1 系统类图", 3)
    add_body(
        doc,
        (
            "从面向对象建模角度看，系统可以抽象为视图层、服务层、控制层和数据实体层四个主要部分。"
            "视图层中的前端视图负责触发页面渲染；接口服务负责向后端发送请求；控制层中的控制器负责"
            "路由分发与结果封装；数据服务层中的数据文件服务和图表服务负责处理结果文件与图表构造；"
            "销售数据实体、模型指标实体与预测结果实体则对应数据库逻辑设计中的核心业务对象。"
        ),
    )
    add_figure(doc, diagrams["class_diagram"], "图3.1 系统类图", width_cm=16)

    add_heading(doc, "3.3.2 系统顺序图", 3)
    add_body(
        doc,
        (
            "顺序图以“查看指定模型预测图”为主要场景。用户首先在页面中选择模型，页面随后请求模"
            "型详情接口获取指标信息，再请求模型图表接口获取折线图HTML。后端在处理该场景时，会先"
            "后调用数据读取服务和图表构建服务，并将结果返回给页面进行展示。顺序图清晰地反映了用户、"
            "页面、API服务、Django控制器和图表服务之间的时序关系。"
        ),
    )
    add_figure(doc, diagrams["sequence_diagram"], "图3.2 系统顺序图", width_cm=16)

    add_heading(doc, "3.3.3 系统活动图", 3)
    add_body(
        doc,
        (
            "活动图描述了用户查看模型预测结果时的流程状态变化，覆盖页面进入、加载模型列表、用户"
            "选择模型、系统调用接口、判断接口是否成功、展示指标与图表以及异常提示等步骤。活动图"
            "有助于论文清楚说明正常流程与异常流程，体现对系统交互过程的完整分析。"
        ),
    )
    add_figure(doc, diagrams["activity_diagram"], "图3.3 系统活动图", width_cm=15.5)

    add_heading(doc, "3.4 数据库设计", 2)
    add_heading(doc, "3.4.1 概念结构设计", 3)
    add_body(
        doc,
        (
            "虽然当前系统运行依赖CSV结果文件，但为满足软件设计类论文对数据库设计的要求，同时为"
            "后续系统工程化扩展提供依据，本文构建了包含系统用户、销售样本、模型信息、模型指标、"
            "预测结果、图表资源和操作日志七类实体的数据库逻辑方案。该方案能够覆盖系统运行所需的"
            "主要业务对象，并支持用户体系、资源管理和日志审计等扩展需求。"
        ),
    )
    add_figure(doc, diagrams["er_diagram"], "图3.4 系统ER图", width_cm=16)

    add_heading(doc, "3.4.2 逻辑结构设计", 3)
    add_body(
        doc,
        (
            "逻辑结构设计中，销售数据表用于存储真实销售数据及其特征字段；模型信息表用于存储模型基"
            "础信息和状态；模型指标表用于存储各模型误差指标；预测结果表用于存储按月份组织的预测值；"
            "图表资源表用于记录图表资源；系统用户表与操作日志表分别用于用户与日志扩展。现阶段虽然"
            "未全部落地为实际数据库表，但其结构设计已经能"
            "够支撑论文中的数据库章节和未来系统扩展。"
        ),
    )
    add_table_caption(doc, "表3.2 用户信息表")
    add_table(
        doc,
        [
            ["字段名", "类型", "主键/外键", "非空", "说明"],
            ["user_id", "INT", "PK", "是", "用户编号"],
            ["username", "VARCHAR(32)", "", "是", "用户名"],
            ["password", "VARCHAR(64)", "", "是", "密码摘要"],
            ["role", "VARCHAR(16)", "", "是", "角色类型"],
            ["create_time", "DATETIME", "", "是", "创建时间"],
        ],
    )
    add_table_caption(doc, "表3.3 销售样本表")
    add_table(
        doc,
        [
            ["字段名", "类型", "主键/外键", "非空", "说明"],
            ["sales_id", "INT", "PK", "是", "销售记录编号"],
            ["sale_month", "VARCHAR(7)", "", "是", "销售月份"],
            ["sales_value", "DECIMAL(10,2)", "", "是", "销量值"],
            ["last_month_sales", "DECIMAL(10,2)", "", "否", "上月销量"],
            ["last_year_same_month", "DECIMAL(10,2)", "", "否", "去年同月销量"],
            ["region", "VARCHAR(32)", "", "是", "销售区域"],
            ["product_series", "VARCHAR(32)", "", "是", "产品系列"],
        ],
    )
    add_table_caption(doc, "表3.4 模型评估表")
    add_table(
        doc,
        [
            ["字段名", "类型", "主键/外键", "非空", "说明"],
            ["metric_id", "INT", "PK", "是", "指标编号"],
            ["model_id", "INT", "FK", "是", "模型编号"],
            ["mae", "DECIMAL(10,4)", "", "是", "平均绝对误差"],
            ["rmse", "DECIMAL(10,4)", "", "是", "均方根误差"],
            ["mape", "DECIMAL(10,4)", "", "是", "平均绝对百分比误差"],
        ],
    )
    add_table_caption(doc, "表3.5 预测结果表")
    add_table(
        doc,
        [
            ["字段名", "类型", "主键/外键", "非空", "说明"],
            ["forecast_id", "INT", "PK", "是", "预测编号"],
            ["model_id", "INT", "FK", "是", "模型编号"],
            ["forecast_month", "VARCHAR(7)", "", "是", "预测月份"],
            ["forecast_value", "DECIMAL(10,4)", "", "是", "预测值"],
        ],
    )

    add_heading(doc, "3.5 安全与性能设计", 2)
    add_body(
        doc,
        (
            "安全设计方面，系统当前通过后端统一读取结果文件和输出图片资源，避免前端直接暴露文件"
            "系统路径，降低资源访问不当带来的风险。若后续增加登录功能，可进一步引入基于Django的"
            "认证机制、权限校验、日志审计和CSRF防护。性能设计方面，系统采用静态结果文件读取与按"
            "需图表渲染相结合的方式，避免每次页面访问都重新训练模型，从而显著降低响应耗时。同时，"
            "前端通过组合式函数集中加载数据，在多数情况下能够减少重复请求，提高页面访问效率。"
        ),
        (
            "由于本系统的主要使用场景是教学展示和论文答辩，因此性能设计强调“稳定可演示”而非“大"
            "规模并发”。不过从架构上看，当前的分层设计已经为后续通过缓存、异步任务和数据库索引等"
            "方式继续优化性能提供了良好基础。"
        ),
    )

    add_heading(doc, "3.6 本章小结", 2)
    add_body(
        doc,
        (
            "本章完成了系统总体架构、前后端模块、接口结构、类图、顺序图、活动图、ER图及数据库"
            "设计的系统说明，形成了完整的软件设计成果。通过上述设计可以看出，系统不仅能够支撑当"
            "前论文所需的展示与分析需求，也具备较好的扩展性和工程演进空间。"
        ),
    )


def chapter_4(doc, stats):
    metrics_df = stats["metrics_df"]
    forecast_df = stats["forecast_df"]

    add_heading(doc, "第4章 系统实现", 1)
    add_heading(doc, "4.1 开发环境与工程结构", 2)
    add_body(
        doc,
        (
            "系统开发环境包括Python 3.13、Django 5.2、Vue 3、Vite 8、Node.js 22、PyTorch和"
            "Pyecharts等组件。就工程结构来看，后端目录主要包含dashboard应用、路由配置、视图文件"
            "和测试文件，前端目录则包含路由配置、公共组件、组合式函数以及各页面视图。页面层、接"
            "口层和测试层拆开之后，整体结构会清楚一些，后面维护起来也更方便，论文整理时也更容易"
            "对应。"
        ),
        (
            "和原来那种单脚本的数据分析方式相比，现在的工程结构已经更接近标准的Web系统项目。前"
            "端负责用户可见交互，后端负责资源组织和统一输出，预测模型结果则作为可复用的数据成果"
            "接入系统。也正因为有了这样的工程结构，论文才真正具备了从“大数据分析类”转向“软件设"
            "计类”的基础。"
        ),
    )

    add_heading(doc, "4.2 前端功能实现", 2)
    add_heading(doc, "4.2.1 系统首页实现", 3)
    add_body(
        doc,
        (
            "系统首页基于两栏式后台布局实现。左侧为固定导航栏，提供系统首页、数据中心、模型评估、"
            "预测结果、图表中心和系统设计六个入口；右侧由顶部标题栏和内容区域组成。首页内容区域"
            "主要展示样本量、预测月份数、模型数量和图表数量等统计卡片，同时以默认模型面板和运行"
            "概览面板集中展示推荐模型、误差指标和预测峰值信息。该页面既可作为系统总入口，也适合"
            "作为论文答辩时的首页展示。"
        ),
        (
            f"当前首页默认展示{stats['best_rmse_name']}为推荐模型，其原因是该模型的RMSE最小，"
            f"仅为{stats['best_rmse']}。通过在首页直接呈现推荐模型、MAE、RMSE和MAPE，系统避免"
            "了用户必须切换多个页面才能获取核心结论的问题，提高了信息传达效率。"
        ),
    )
    add_figure(doc, SYSTEM_SCREENSHOTS["dashboard"], "图4.1 系统首页运行界面", width_cm=16)
    add_heading(doc, "4.2.2 数据中心实现", 3)
    add_body(
        doc,
        (
            "数据中心页面通过样本数据接口读取真实销售记录，并采用通用DataTable组件统一展示各字段。"
            "页面顶部以统计卡片展示样本条数、起始月份、结束月份和模型数量，主区域则展示技术栈标"
            "签和原始月度数据预览。页面所展示的数据来源于真实销售出库整理结果，"
            "从而保证系统不是空壳演示，而是建立在实际业务样本之上。"
        ),
        (
            "DataTable组件通过传入列名数组和行数据数组即可完成不同页面的表格展示，这种组件复用"
            "方式显著减少了重复代码，也使模型评估页和预测结果页可以沿用同一套表格样式，体现了前"
            "端模块化设计思想。"
        ),
    )
    add_figure(doc, SYSTEM_SCREENSHOTS["data_center"], "图4.2 数据中心运行界面", width_cm=16)

    add_heading(doc, "4.3 后端接口实现", 2)
    add_heading(doc, "4.3.1 结果聚合接口实现", 3)
    add_body(
        doc,
        (
            "后端视图层通过overview、metrics、forecast、sample_data、modules、insights和"
            "model_detail等函数分别对不同业务能力进行封装。其中overview接口负责返回项目基础概"
            "况，metrics接口负责读取模型评估CSV并确定推荐模型，forecast接口负责返回未来预测结果，"
            "sample_data接口负责截取月度样本前若干条数据进行展示，model_detail接口则用于根据模型"
            "名提取指标和时间序列。这些接口共同构成系统的数据服务基础。"
        ),
        (
            "实现过程中，后端统一通过read_csv_or_empty函数读取文件，并通过records函数将DataFrame"
            "转换为适合JSON输出的结构，确保不同接口在空值处理、字段输出和数据格式上保持一致。对"
            "前端来说，这种封装显著降低了页面编写复杂度。"
        ),
    )
    add_figure(doc, SYSTEM_SCREENSHOTS["system"], "图4.3 系统设计页面运行界面", width_cm=16)
    add_heading(doc, "4.3.2 图表接口实现", 3)
    add_body(
        doc,
        (
            "图表接口通过build_metric_chart、build_multi_forecast_chart、build_model_forecast_chart"
            "和build_sales_trend_chart等函数构建不同类型的Pyecharts图表。系统将图表最终渲染为"
            "独立HTML页面，并通过iframe嵌入前端页面中。为保证同源嵌入正常显示，后端配置了相应的"
            "iframe访问策略。该实现方式使前端无需直接操控图表配置，也使后端能统一管理图表样式和"
            "标题。"
        ),
        (
            "在实现过程中，本文将流程图、用例图、类图、顺序图、活动图和ER图统一采用Mermaid建"
            "模后导出为图片，直接插入到系统分析与系统设计章节中；而系统运行界面则通过浏览器自动"
            "化方式生成截图插入到第四章。这样的处理方式既保证了图形内容与正文的一致性，也减少了"
            "人工重复绘图带来的版式误差。"
        ),
    )

    add_heading(doc, "4.4 预测与可视化实现", 2)
    add_heading(doc, "4.4.1 模型评估模块实现", 3)
    add_body(
        doc,
        (
            "模型评估模块的实现包括评估表渲染与图表联动两个部分。页面首次加载后，通过metrics接口"
            "获取全部模型的MAE、RMSE和MAPE值，并据此计算RMSE最优模型与MAPE最优模型。页面上方以"
            "统计卡片形式分别显示这两个结果，下方则使用高亮表格显示全部指标。这样一来，用户既可"
            "以一眼看到推荐模型，也可以从整体表格中比较不同模型之间的差异。"
        ),
    )
    add_table_caption(doc, "表4.1 模型评估结果展示表")
    rows = [["模型", "MAE", "RMSE", "MAPE(%)"]]
    for _, row in metrics_df.iterrows():
        rows.append(
            [
                row["模型"],
                f"{row['MAE']:.4f}",
                f"{row['RMSE']:.4f}",
                f"{row['MAPE(%)']:.4f}",
            ]
        )
    add_table(doc, rows)
    add_figure(doc, SYSTEM_SCREENSHOTS["metrics"], "图4.4 模型评估页面运行界面", width_cm=16)
    add_figure(doc, CHART_IMAGES["mae"], "图4.5 MAE指标对比图", width_cm=14.5)
    add_figure(doc, CHART_IMAGES["rmse"], "图4.6 RMSE指标对比图", width_cm=14.5)
    add_figure(doc, CHART_IMAGES["mape"], "图4.7 MAPE指标对比图", width_cm=14.5)

    add_heading(doc, "4.4.2 预测结果模块实现", 3)
    add_body(
        doc,
        (
            "预测结果模块采用“筛选条件 + 重点模型详情 + 全量预测表”组合布局。用户可以通过月份关键"
            "字筛选预测表记录，也可以在下拉框中选择重点模型查看其MAE、RMSE、MAPE和局部预测值。"
            "当用户切换模型时，页面通过watch监听自动重新请求模型详情接口，从而保证指标信息与局部"
            "预览同步更新。该设计增强了页面的交互性，使预测结果展示不再停留于静态表格。"
        ),
        (
            f"从当前结果看，未来12个月预测峰值出现在{stats['peak_month']}，对应模型为"
            f"{stats['peak_model']}，预测值为{stats['peak_value']}。系统通过重点模型预览和整表"
            "展示的双层结构，使用户既能快速捕捉关键节点，也能保留整体比较能力。"
        ),
    )
    add_table_caption(doc, "表4.2 部分未来预测结果展示表")
    preview_rows = [forecast_df.columns.tolist()]
    for _, row in forecast_df.head(6).iterrows():
        preview_rows.append([str(item) for item in row.tolist()])
    add_table(doc, preview_rows)
    add_figure(doc, SYSTEM_SCREENSHOTS["forecast"], "图4.8 预测结果页面运行界面", width_cm=16)

    add_heading(doc, "4.4.3 图表中心实现", 3)
    add_body(
        doc,
        (
            "图表中心页面根据查询参数和模型选择结果动态组织图表列表。当用户选择“指标图”类别时，"
            "系统展示MAE、RMSE和MAPE三张独立柱状图；当用户选择“趋势图”类别时，系统展示历史销量"
            "趋势图和多模型未来预测图；无论在哪种类别下，系统都会根据当前所选模型追加该模型的专"
            "属预测曲线图。通过这种方式，图表中心既支持按类别查看，也支持按模型深入分析。"
        ),
        (
            "图表中心是系统中最能体现软件化成果的模块。与传统论文中直接插入静态图不同，该模块"
            "允许在页面内切换模型并即时刷新图形结果，使图表真正成为可交互的系统功能。"
        ),
    )
    add_figure(doc, SYSTEM_SCREENSHOTS["charts"], "图4.9 图表中心运行界面", width_cm=16)
    add_figure(doc, CHART_IMAGES["multi"], "图4.10 未来12个月多模型预测图", width_cm=15.5)
    add_figure(doc, CHART_IMAGES["trend"], "图4.11 历史与未来销量趋势图", width_cm=15.5)

    add_heading(doc, "4.4.4 LSTM模型集成实现", 3)
    add_body(
        doc,
        (
            "为满足论文中多模型比较和深度学习扩展的要求，系统在保留传统机器学习模型的同时，引入"
            "了基于PyTorch实现的LSTM模型。LSTM通过时间序列窗口化数据输入完成训练，其结果与其他"
            "模型一样被统一整理为评估文件和预测文件，并通过相同的详情接口和图表接口进行展示。由"
            "于系统层面对模型采用统一服务接口，因此LSTM的接入并未破坏已有页面结构。"
        ),
        (
            "从结果来看，LSTM在当前小样本数据集上的RMSE并非最优，但其引入价值在于展示了系统具备"
            "兼容传统机器学习和深度学习模型的能力。对于软件设计类论文而言，这种“算法可插拔”的设"
            "计更能体现系统扩展性。"
        ),
    )
    add_figure(doc, CHART_IMAGES["loss"], "图4.12 LSTM训练损失图", width_cm=14.5)

    add_heading(doc, "4.5 本章小结", 2)
    add_body(
        doc,
        (
            "本章从前端页面、后端接口、预测结果和可视化模块四个方面说明了系统实现过程。实践表明，"
            "该系统已经具备真实数据展示、多模型切换、动态图表渲染和模块化组织等完整软件能力，能"
            "够支撑软件设计类毕业论文对“实现部分”的要求。"
        ),
    )


def chapter_5(doc):
    add_heading(doc, "第5章 系统测试", 1)
    add_heading(doc, "5.1 测试环境与策略", 2)
    add_heading(doc, "5.1.1 测试环境", 3)
    add_body(
        doc,
        (
            "系统测试环境为macOS开发环境。后端在Python 3.13和Django 5.2下运行，前端在Node.js 22、"
            "Vue 3和Vite环境下运行。测试阶段采用本地前后端分离方式，前端开发服务器通过代理访问后"
            "端接口，后端返回JSON数据和Pyecharts图表页面。测试时，前端使用Vitest对路由、接口封装"
            "和共享数据加载逻辑进行单元测试，后端使用Django TestCase验证接口逻辑，同时结合页面功"
            "能测试表检查各模块在演示场景下的可用性。"
        ),
    )
    add_heading(doc, "5.1.2 测试策略", 3)
    add_body(
        doc,
        (
            "测试策略分为模块单元测试、接口联调测试和运行结果分析三个层次。模块单元测试重点验证"
            "首页、数据中心、模型评估、预测结果和图表中心等模块是否能够正确调用对应接口并完成展"
            "示；接口联调测试重点验证概览、评估、预测、详情、图表和图片接口返回是否正确；运行结"
            "果分析则从功能完整性、页面稳定性和图表联动性三个方面综合评价系统效果。"
        ),
    )

    add_heading(doc, "5.2 模块单元测试", 2)
    add_heading(doc, "5.2.1 系统首页模块测试", 3)
    add_body(
        doc,
        (
            "系统首页模块重点测试概览数据加载、推荐模型展示和快捷入口跳转功能。测试结果表明，首"
            "页能够正常展示统计卡片与推荐模型，且跳转到模型评估、预测结果、图表中心和系统设计页"
            "面的路由均可正常工作。"
        ),
    )
    add_table_caption(doc, "表5.1 系统首页模块测试结果")
    add_table(
        doc,
        [
            ["测试编号", "测试内容", "预期结果", "实际结果", "结论"],
            ["UT-01", "首页概览加载", "统计卡片正常显示", "通过", "通过"],
            ["UT-02", "推荐模型展示", "首页显示最佳RMSE模型", "通过", "通过"],
            ["UT-03", "快捷入口跳转", "能够跳转至对应页面", "通过", "通过"],
        ],
    )

    add_heading(doc, "5.2.2 数据中心模块测试", 3)
    add_body(
        doc,
        (
            "数据中心模块重点测试真实样本数据加载、时间范围展示与表格渲染。测试表明，页面能够正"
            "确加载样本数据接口并展示字段内容，表格列数与数据行内容与原始样本保持一致。"
        ),
    )
    add_table_caption(doc, "表5.2 数据中心模块测试结果")
    add_table(
        doc,
        [
            ["测试编号", "测试内容", "预期结果", "实际结果", "结论"],
            ["UT-04", "样本数据接口调用", "返回样本列和数据项", "通过", "通过"],
            ["UT-05", "时间范围显示", "页面显示起止月份", "通过", "通过"],
            ["UT-06", "表格组件渲染", "列名与数据对应正确", "通过", "通过"],
        ],
    )

    add_heading(doc, "5.2.3 模型评估模块测试", 3)
    add_body(
        doc,
        (
            "模型评估模块重点测试指标列表加载、推荐模型高亮和图表入口跳转。测试结果显示，该页面"
            "能够正确加载评估结果、计算最佳模型，并支持跳转至图表中心查看指标图。"
        ),
    )
    add_table_caption(doc, "表5.3 模型评估模块测试结果")
    add_table(
        doc,
        [
            ["测试编号", "测试内容", "预期结果", "实际结果", "结论"],
            ["UT-07", "指标表加载", "显示全部模型指标", "通过", "通过"],
            ["UT-08", "推荐模型高亮", "最佳模型高亮显示", "通过", "通过"],
            ["UT-09", "图表入口跳转", "跳转至指标图页面", "通过", "通过"],
        ],
    )

    add_heading(doc, "5.2.4 预测结果模块测试", 3)
    add_body(
        doc,
        (
            "预测结果模块重点测试月份筛选、重点模型切换和详情联动。测试表明，用户输入月份关键字"
            "后能够正确过滤表格记录，切换重点模型时页面会同步刷新指标信息和局部预测预览。"
        ),
    )
    add_table_caption(doc, "表5.4 预测结果模块测试结果")
    add_table(
        doc,
        [
            ["测试编号", "测试内容", "预期结果", "实际结果", "结论"],
            ["UT-10", "月份筛选", "按关键字过滤预测表", "通过", "通过"],
            ["UT-11", "重点模型切换", "详情数据同步刷新", "通过", "通过"],
            ["UT-12", "局部预览展示", "显示所选模型最近预览值", "通过", "通过"],
        ],
    )

    add_heading(doc, "5.2.5 图表中心模块测试", 3)
    add_body(
        doc,
        (
            "图表中心模块重点测试图表分类加载、模型预测图切换和iframe嵌入显示。测试结果表明，"
            "系统能够在页面中正常嵌入Pyecharts图表，并在切换模型后刷新指定模型预测图。"
        ),
    )
    add_table_caption(doc, "表5.5 图表中心模块测试结果")
    add_table(
        doc,
        [
            ["测试编号", "测试内容", "预期结果", "实际结果", "结论"],
            ["UT-13", "指标图展示", "MAE/RMSE/MAPE图正常显示", "通过", "通过"],
            ["UT-14", "趋势图展示", "历史趋势图与多模型图正常显示", "通过", "通过"],
            ["UT-15", "模型图切换", "切换模型后图表刷新", "通过", "通过"],
        ],
    )

    add_heading(doc, "5.2.6 后端接口模块测试", 3)
    add_body(
        doc,
        (
            "后端接口模块采用Django TestCase进行自动化单元测试。测试覆盖overview、sample-data、"
            "metrics、forecast、model-detail、modules、insights以及图表和图片资源等多个接口。"
            "与此同时，前端还使用Vitest对路由定义、API请求路径拼装和useSystemData共享加载逻辑进"
            "行测试。通过对状态码、返回字段、URL编码和共享状态赋值进行断言，验证了前后端核心模块"
            "的可用性。"
        ),
    )
    add_table_caption(doc, "表5.6 自动化单元测试结果")
    add_table(
        doc,
        [
            ["测试编号", "接口或模块", "断言内容", "结果"],
            ["AT-00", "前端路由测试", "首页、数据中心、模型评估等路由均存在", "通过"],
            ["AT-00-2", "前端共享数据测试", "共享状态能正确写入概览、预测和模块数据", "通过"],
            ["AT-01", "/api/overview/", "返回200且包含project_name", "通过"],
            ["AT-02", "/api/sample-data/", "返回200且包含items", "通过"],
            ["AT-03", "/api/metrics/", "返回200且包含best_model", "通过"],
            ["AT-04", "/api/forecast/", "返回200且包含columns和items", "通过"],
            ["AT-05", "/api/model-detail/", "返回指定模型详情", "通过"],
            ["AT-06", "/api/charts/model-forecast/", "返回包含echarts的HTML", "通过"],
        ],
    )

    add_heading(doc, "5.3 接口联调与结果分析", 2)
    add_body(
        doc,
        (
            "在前后端联调阶段，系统概览页、模型评估页、预测结果页和图表中心页均能够通过统一数据"
            "加载逻辑获取所需数据。由于后端接口统一遵循JSON结构，前端在处理概览、表格和详情时不"
            "需要对不同接口单独适配，从而降低了联调复杂度。图表接口则通过HTML形式返回，前端只需"
            "构造iframe地址即可完成嵌入展示。联调结果说明当前接口设计与页面设计之间具有较好匹配"
            "性。"
        ),
        (
            "综合测试结果可以看出，系统已经具备较强的展示稳定性和模块协同能力。首页能够快速给出"
            "推荐模型和系统规模；数据中心能够证明数据来源真实；模型评估模块能够展示多模型对比结"
            "果；预测结果模块能够支持条件筛选和模型详情查看；图表中心能够完成动态图表联动。这些"
            "结果说明系统不仅满足论文写作需要，也具备一定实际使用价值。"
        ),
        (
            "从论文表达效果来看，系统截图、动态图表和Mermaid建模图共同构成了较完整的实现证据链。"
            "其中，建模图负责说明系统结构与业务流程，页面截图负责说明功能落地情况，测试表负责说"
            "明系统运行结果，三者相互支撑，使论文主体内容更加完整。"
        ),
    )
    add_table_caption(doc, "表5.7 核心接口联调测试结果")
    add_table(
        doc,
        [
            ["接口路径", "测试目标", "测试结果"],
            ["/api/overview/", "概览数据与首页绑定", "通过"],
            ["/api/metrics/", "模型评估表与推荐模型绑定", "通过"],
            ["/api/forecast/", "预测结果表格展示", "通过"],
            ["/api/model-detail/", "重点模型详情刷新", "通过"],
            ["/api/charts/rmse/", "RMSE图表嵌入", "通过"],
            ["/api/charts/model-forecast/", "指定模型预测图刷新", "通过"],
        ],
    )

    add_heading(doc, "5.4 本章小结", 2)
    add_body(
        doc,
        (
            "本章通过模块单元测试、接口自动化测试和联调测试验证了系统的主要功能和运行稳定性。"
            "测试结果表明，系统在页面跳转、数据加载、模型切换和动态图表展示方面均达到预期目标，"
            "能够满足软件设计类毕业论文对测试章节和系统可运行性的要求。"
        ),
    )


def chapter_6(doc):
    add_heading(doc, "第6章 总结与展望", 1)
    add_heading(doc, "6.1 总结", 2)
    add_body(
        doc,
        (
            "本文围绕冰柠代理销售业务，完成了基于Vue与Django的销售分析与预测系统设计与实"
            "现。和原来偏重数据分析的论文写法相比，这次把真实销售数据、多模型评估、未来预测和可"
            "视化展示放到了同一个可运行的软件平台里，同时补充完成了软件设计类毕业论文通常需要的"
            "业务流程、用例分析、类图、顺序图、活动图、ER图、数据库设计和模块测试等内容。"
        ),
        (
            "从实现情况看，前端已经完成了系统首页、数据中心、模型评估、预测结果、图表中心和系统"
            "设计等页面；后端也完成了概览、样本、评估、预测、模型详情、图表、图片资源和模块说明"
            "等接口；模型层接入了随机森林、梯度提升回归、极端随机森林、线性回归和LSTM；图表层则"
            "借助Pyecharts实现了指标图和趋势图的动态展示。测试结果说明，系统能够稳定运行，扩展性"
            "也比较符合预期。"
        ),
    )
    add_heading(doc, "6.2 展望", 2)
    add_body(
        doc,
        (
            "后续工作可从三个方向继续完善。第一，增强业务完整性，增加登录认证、角色权限、操作日"
            "志和异常告警等功能，使系统更加接近真实企业应用环境。第二，增强数据管理能力，将当前"
            "CSV结果文件逐步迁移至关系数据库，并支持在线数据上传、模型重训和结果归档。第三，增"
            "强分析深度，在保留当前多模型对比的基础上，引入更细粒度的特征分析、区域分组预测和模"
            "型解释能力，从而进一步提升系统的决策支持价值。"
        ),
    )


def references_and_appendix(doc):
    add_heading(doc, "参考文献", 1)
    for ref in REFERENCES:
        add_paragraph(doc, ref, style="Normal", role="body")

    add_heading(doc, "致谢", 1)
    add_body(
        doc,
        (
            "感谢指导教师在课题改题、系统设计、论文结构调整和测试完善过程中给予的指导与帮助。"
            "感谢同学在系统演示、素材整理和格式校对中的支持。"
        ),
    )


def build_document():
    stats = build_stats()
    diagrams = render_mermaid_diagrams()
    doc = Document(SOURCE)
    update_front_matter(doc)
    remove_body_from_paragraph(doc, 84)
    chapter_1(doc, stats)
    doc.add_page_break()
    chapter_2(doc, stats, diagrams)
    doc.add_page_break()
    chapter_3(doc, stats, diagrams)
    doc.add_page_break()
    chapter_4(doc, stats)
    doc.add_page_break()
    chapter_5(doc)
    doc.add_page_break()
    chapter_6(doc)
    doc.add_page_break()
    references_and_appendix(doc)
    normalize_caption_numbers(doc)
    doc.save(TARGET)


if __name__ == "__main__":
    build_document()
